import os
import git
from pathlib import Path
from tqdm import tqdm
from typing import List, Dict, Any
import mimetypes

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate

# File parsing libraries
import csv
import json
import xml.etree.ElementTree as ET
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    import openpyxl
except ImportError:
    openpyxl = None
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
try:
    import markdown
except ImportError:
    markdown = None

# ========== CONFIG ==========
REPO_PATH = "./target_repo"  # Change to your repo path
PERSIST_DIR = "./data/repo_index"
EMBEDDING_MODEL = "nomic-embed-text"
LLM_MODEL = "codellama:latest"
REFLECTION_MODEL = "llama3.2:latest"  # Secondary model for reflection
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
USE_REFLECTION = True  # Toggle reflection layer

# File extensions to parse
CODE_EXTENSIONS = {'.py', '.js', '.java', '.cpp', '.c', '.h', '.cs', '.go', '.rs', '.php', '.rb', '.swift', '.kt'}
TEXT_EXTENSIONS = {'.txt', '.md', '.rst', '.log', '.yaml', '.yml', '.json', '.xml', '.html', '.css', '.env'}
DOCUMENT_EXTENSIONS = {'.pdf', '.docx', '.doc', '.csv', '.xlsx', '.xls'}

# ========== FILE PARSERS ==========

def parse_text_file(file_path: str) -> str:
    """Parse plain text files"""
    encodings = ['utf-8', 'latin-1', 'cp1252']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return f"[Unable to decode file: {file_path}]"

def parse_csv_file(file_path: str) -> str:
    """Parse CSV files"""
    try:
        with open(file_path, 'r', encoding='utf-8', newline='') as f:
            reader = csv.reader(f)
            rows = list(reader)
            if not rows:
                return ""
            # Create a formatted string representation
            header = rows[0] if rows else []
            content = f"CSV File: {os.path.basename(file_path)}\n"
            content += f"Columns: {', '.join(header)}\n"
            content += f"Rows: {len(rows) - 1}\n\n"
            # Include first few rows as sample
            content += "Sample data:\n"
            for row in rows[:min(10, len(rows))]:
                content += " | ".join(str(cell) for cell in row) + "\n"
            return content
    except Exception as e:
        return f"[Error parsing CSV {file_path}: {str(e)}]"

def parse_json_file(file_path: str) -> str:
    """Parse JSON files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return f"JSON File: {os.path.basename(file_path)}\n{json.dumps(data, indent=2)}"
    except Exception as e:
        return f"[Error parsing JSON {file_path}: {str(e)}]"

def parse_xml_file(file_path: str) -> str:
    """Parse XML files"""
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        content = f"XML File: {os.path.basename(file_path)}\n"
        content += f"Root element: {root.tag}\n"
        content += ET.tostring(root, encoding='unicode', method='xml')
        return content
    except Exception as e:
        return f"[Error parsing XML {file_path}: {str(e)}]"

def parse_pdf_file(file_path: str) -> str:
    """Parse PDF files"""
    if PyPDF2 is None:
        return f"[PDF parsing unavailable - install PyPDF2: pip install PyPDF2]"
    
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            content = f"PDF File: {os.path.basename(file_path)}\n"
            content += f"Pages: {len(reader.pages)}\n\n"
            for i, page in enumerate(reader.pages[:20]):  # Limit to first 20 pages
                text = page.extract_text()
                if text.strip():
                    content += f"--- Page {i+1} ---\n{text}\n"
            return content
    except Exception as e:
        return f"[Error parsing PDF {file_path}: {str(e)}]"

def parse_docx_file(file_path: str) -> str:
    """Parse DOCX files"""
    if DocxDocument is None:
        return f"[DOCX parsing unavailable - install python-docx: pip install python-docx]"
    
    try:
        doc = DocxDocument(file_path)
        content = f"DOCX File: {os.path.basename(file_path)}\n\n"
        for para in doc.paragraphs:
            if para.text.strip():
                content += para.text + "\n"
        return content
    except Exception as e:
        return f"[Error parsing DOCX {file_path}: {str(e)}]"

def parse_excel_file(file_path: str) -> str:
    """Parse Excel files"""
    if openpyxl is None:
        return f"[Excel parsing unavailable - install openpyxl: pip install openpyxl]"
    
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True)
        content = f"Excel File: {os.path.basename(file_path)}\n"
        content += f"Sheets: {', '.join(wb.sheetnames)}\n\n"
        
        for sheet_name in wb.sheetnames[:5]:  # Limit to first 5 sheets
            sheet = wb[sheet_name]
            content += f"\n--- Sheet: {sheet_name} ---\n"
            rows = list(sheet.iter_rows(values_only=True, max_row=20))
            for row in rows:
                content += " | ".join(str(cell) if cell is not None else "" for cell in row) + "\n"
        
        return content
    except Exception as e:
        return f"[Error parsing Excel {file_path}: {str(e)}]"

def parse_markdown_file(file_path: str) -> str:
    """Parse Markdown files"""
    try:
        content = parse_text_file(file_path)
        if markdown:
            # Convert to HTML for better structure extraction
            html = markdown.markdown(content)
            return f"Markdown File: {os.path.basename(file_path)}\n\n{content}"
        return content
    except Exception as e:
        return f"[Error parsing Markdown {file_path}: {str(e)}]"

def parse_file(file_path: str) -> str:
    """Route to appropriate parser based on file extension"""
    ext = Path(file_path).suffix.lower()
    
    # CSV
    if ext == '.csv':
        return parse_csv_file(file_path)
    
    # JSON
    elif ext == '.json':
        return parse_json_file(file_path)
    
    # XML
    elif ext == '.xml':
        return parse_xml_file(file_path)
    
    # PDF
    elif ext == '.pdf':
        return parse_pdf_file(file_path)
    
    # DOCX
    elif ext in {'.docx', '.doc'}:
        return parse_docx_file(file_path)
    
    # Excel
    elif ext in {'.xlsx', '.xls'}:
        return parse_excel_file(file_path)
    
    # Markdown
    elif ext == '.md':
        return parse_markdown_file(file_path)
    
    # Default: plain text
    else:
        return parse_text_file(file_path)

# ========== LOADING FUNCTIONS ==========

def load_code_files(repo_path: str) -> List[Document]:
    """Load all supported file types from repository"""
    documents = []
    repo_path_obj = Path(repo_path)
    
    all_extensions = CODE_EXTENSIONS | TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS
    
    for ext in all_extensions:
        files = list(repo_path_obj.rglob(f"*{ext}"))
        for file_path in tqdm(files, desc=f"Loading {ext} files"):
            # Skip hidden files and common ignore patterns
            if any(part.startswith('.') for part in file_path.parts):
                continue
            if any(skip in str(file_path) for skip in ['node_modules', '__pycache__', '.git', 'venv', '.venv']):
                continue
            
            try:
                content = parse_file(str(file_path))
                if content and len(content.strip()) > 0:
                    rel_path = file_path.relative_to(repo_path_obj)
                    documents.append(
                        Document(
                            page_content=content,
                            metadata={
                                "source": str(rel_path),
                                "type": "file",
                                "extension": ext,
                                "full_path": str(file_path)
                            }
                        )
                    )
            except Exception as e:
                print(f"⚠️ Error loading {file_path}: {e}")
    
    print(f"✅ Loaded {len(documents)} files")
    return documents

def load_git_commits(repo_path: str, branch: str = "main", max_commits: int = 100) -> List[Document]:
    """Load recent Git commit messages and diffs"""
    try:
        repo = git.Repo(repo_path)
        commits = list(repo.iter_commits(branch, max_count=max_commits))
        
        documents = []
        for commit in tqdm(commits, desc="Loading commits"):
            message = f"Commit: {commit.hexsha[:7]}\n"
            message += f"Author: {commit.author.name}\n"
            message += f"Date: {commit.committed_datetime}\n"
            message += f"Message: {commit.message}\n"
            
            # Add diff info
            if commit.parents:
                diffs = commit.parents[0].diff(commit, create_patch=True)
                message += "\nChanges:\n"
                for diff in diffs[:10]:  # Limit diffs
                    message += f"- {diff.a_path}\n"
            
            documents.append(
                Document(
                    page_content=message,
                    metadata={
                        "source": f"commit-{commit.hexsha[:7]}",
                        "type": "commit",
                        "author": commit.author.name,
                        "date": str(commit.committed_datetime)
                    }
                )
            )
        
        print(f"✅ Loaded {len(documents)} commits")
        return documents
    except Exception as e:
        print(f"⚠️ Could not load commits: {e}")
        return []

# ========== INDEXING ==========

def build_vector_store(documents: List[Document]) -> Chroma:
    """Build or load Chroma vector store"""
    embeddings = OllamaEmbeddings(model=EMBEDDING_MODEL)
    
    if os.path.exists(PERSIST_DIR):
        print("📂 Loading existing index...")
        vectorstore = Chroma(
            persist_directory=PERSIST_DIR,
            embedding_function=embeddings
        )
    else:
        print("🔨 Building new index...")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP
        )
        
        split_docs = text_splitter.split_documents(documents)
        print(f"📄 Created {len(split_docs)} chunks")
        
        vectorstore = Chroma.from_documents(
            documents=split_docs,
            embedding=embeddings,
            persist_directory=PERSIST_DIR
        )
        print("✅ Index built and persisted")
    
    return vectorstore

# ========== REFLECTION LAYER ==========

def create_reflection_prompt() -> PromptTemplate:
    """Create prompt template for reflection layer"""
    template = """You are a refinement assistant. Your job is to improve and clarify an initial answer.

Original Question: {question}

Initial Answer: {initial_answer}

Source Documents Referenced:
{sources}

Please:
1. Verify the accuracy and completeness of the initial answer
2. Add any missing important details from the sources
3. Improve clarity and structure
4. Correct any inconsistencies or errors
5. Provide a more comprehensive and well-organized response

Refined Answer:"""
    
    return PromptTemplate(
        input_variables=["question", "initial_answer", "sources"],
        template=template
    )

def reflect_on_answer(question: str, initial_answer: str, source_docs: List[Document]) -> str:
    """Use reflection model to refine the initial answer"""
    try:
        print("🤔 Reflecting on answer...")
        
        # Prepare source context
        sources_text = "\n\n".join([
            f"Source: {doc.metadata.get('source', 'unknown')}\nContent: {doc.page_content[:500]}..."
            for doc in source_docs[:3]
        ])
        
        # Create reflection chain
        reflection_llm = ChatOllama(model=REFLECTION_MODEL, temperature=0.3)
        prompt = create_reflection_prompt()
        
        # Generate refined answer
        refined = reflection_llm.invoke(
            prompt.format(
                question=question,
                initial_answer=initial_answer,
                sources=sources_text
            )
        )
        
        return refined.content
    except Exception as e:
        print(f"⚠️ Reflection failed: {e}")
        return initial_answer

# ========== Q&A ==========

def create_qa_chain(vectorstore: Chroma):
    """Create RAG Q&A chain"""
    llm = ChatOllama(model=LLM_MODEL, temperature=0)
    
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vectorstore.as_retriever(search_kwargs={"k": 5}),
        return_source_documents=True
    )
    
    return qa_chain

def answer_question(qa_chain, query: str, use_reflection: bool = True) -> Dict[str, Any]:
    """Answer a question with optional reflection layer"""
    # Get initial answer from RAG
    print("🔍 Searching knowledge base...")
    result = qa_chain.invoke({"query": query})
    
    initial_answer = result['result']
    source_docs = result.get('source_documents', [])
    
    # Apply reflection if enabled
    if use_reflection and USE_REFLECTION:
        refined_answer = reflect_on_answer(query, initial_answer, source_docs)
        return {
            'initial_answer': initial_answer,
            'refined_answer': refined_answer,
            'sources': source_docs,
            'used_reflection': True
        }
    else:
        return {
            'initial_answer': initial_answer,
            'refined_answer': initial_answer,
            'sources': source_docs,
            'used_reflection': False
        }

def main():
    print("🚀 Local Code Q&A Assistant with Reflection Layer\n")
    
    # Load documents
    print("📥 Loading repository content...")
    code_docs = load_code_files(REPO_PATH)
    
    # Check if Git repo exists before loading commits
    if os.path.exists(os.path.join(REPO_PATH, '.git')):
        commit_docs = load_git_commits(REPO_PATH, branch="main")
    else:
        print("ℹ️  No Git repository detected, skipping commit history")
        commit_docs = []
    
    all_docs = code_docs + commit_docs
    
    if not all_docs:
        print("❌ No documents found!")
        return
    
    # Build vector store
    vectorstore = build_vector_store(all_docs)
    
    # Create Q&A chain
    qa_chain = create_qa_chain(vectorstore)
    
    # Show reflection status
    reflection_status = "✅ ENABLED" if USE_REFLECTION else "❌ DISABLED"
    print(f"\n🔄 Reflection Layer: {reflection_status}")
    print(f"🤖 Primary Model: {LLM_MODEL}")
    if USE_REFLECTION:
        print(f"🧠 Reflection Model: {REFLECTION_MODEL}")
    
    # Interactive loop
    print("\n💬 Ask questions about the repository (type 'exit' to quit)\n")
    print("Commands:")
    print("  - Type your question normally")
    print("  - 'toggle reflection' - Enable/disable reflection layer")
    print("  - 'exit' - Quit\n")
    
    global USE_REFLECTION
    
    while True:
        query = input("🔍 Query> ").strip()
        
        if query.lower() in ['exit', 'quit', 'q']:
            print("👋 Goodbye!")
            break
        
        if query.lower() == 'toggle reflection':
            USE_REFLECTION = not USE_REFLECTION
            status = "enabled" if USE_REFLECTION else "disabled"
            print(f"🔄 Reflection layer {status}\n")
            continue
        
        if not query:
            continue
        
        # Get answer with optional reflection
        result = answer_question(qa_chain, query, use_reflection=USE_REFLECTION)
        
        if result['used_reflection']:
            print(f"\n💭 Initial Answer:\n{result['initial_answer']}\n")
            print(f"✨ Refined Answer:\n{result['refined_answer']}\n")
        else:
            print(f"\n💡 Answer:\n{result['refined_answer']}\n")
        
        if result['sources']:
            print("📚 Sources:")
            for doc in result['sources'][:3]:
                print(f" - {doc.metadata.get('source', 'unknown')}")
        
        print()

if __name__ == "__main__":
    main()
